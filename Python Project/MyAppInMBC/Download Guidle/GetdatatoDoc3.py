import os
import time
import json
import requests
import sys
from urllib.parse import urljoin, urlparse
import re
from bs4 import BeautifulSoup
from tkinter import Tk, filedialog, messagebox, StringVar, Entry, Button, Label, Listbox, END, Frame, Toplevel, Text
from tkinter.scrolledtext import ScrolledText
from tkinter import Radiobutton, Checkbutton, BooleanVar
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from pathlib import Path
from PIL import Image
from io import BytesIO
import shutil

CONFIG_FILE = "config.json"
CONFIG_DIR = "configs"
DATA_DIR = os.path.join(os.path.dirname(__file__), "data")
SAVED_LINKS_DIR = os.path.join(DATA_DIR, "saved_links")
IMAGE_TMP_DIR = os.path.join(DATA_DIR, "images")
COMPLETED_LINKS_FILE = os.path.join(DATA_DIR, "completed_links.json")

def default_config():
    return {
        "config_name": "Cau hinh mac dinh",
        "base_url": "https://vi.extendoffice.com",
        "menu_selector": "ul#ul-search a",
        "menu_selector_type": "css",
        "ignore_selectors": [
            ".uk-margin-remove-last-child.custom",
            "div.uk-margin-remove-last-child.custom style",
            ".related-articles",
            ".social-share"
        ],
        "ignore_selectors_type": "css",
        "output_docx": "output.docx",
        "link_type": "absolute",
        "relative_base_url": "",
        "content_selectors": [
            "article",
            "main",
            "div[class*='content']",
            "div[class*='article']",
            "div[class*='post']",
            ".entry-content",
            ".post-content"
        ]
    }

if not os.path.exists(CONFIG_DIR):
    os.makedirs(CONFIG_DIR)
if not os.path.exists(DATA_DIR):
    os.makedirs(DATA_DIR)
if not os.path.exists(SAVED_LINKS_DIR):
    os.makedirs(SAVED_LINKS_DIR)
if not os.path.exists(IMAGE_TMP_DIR):
    os.makedirs(IMAGE_TMP_DIR)

def save_config(config, name=None):
    if name:
        path = os.path.join(CONFIG_DIR, f"{name}.json")
    else:
        path = CONFIG_FILE
    with open(path, "w", encoding="utf-8") as f:
        json.dump(config, f, ensure_ascii=False, indent=2)

def load_config(name=None):
    if name:
        path = os.path.join(CONFIG_DIR, f"{name}.json")
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return default_config()

def log_to_gui(log_widget, msg):
    if not log_widget:
        print(msg)
        return
    log_widget.config(state="normal")
    log_widget.insert(END, msg + "\n")
    log_widget.see(END)
    log_widget.config(state="disabled")

def get_links_from_page_css(url, selector):
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
        resp = requests.get(url, headers=headers, timeout=15)
        soup = BeautifulSoup(resp.content, 'html.parser')
        items = soup.select(selector)
        results = []
        for i in items:
            text = i.get_text(strip=True)
            href = i.get('href') or i.get('data-href')
            if href and href.strip().startswith('javascript:'):
                href = None
            if href and not href.startswith('http') and not href.startswith('/'):
                href = '/' + href
            results.append((text, href))
        return results
    except Exception as e:
        print(f"Lỗi get_links_from_page_css: {e}")
        return []

def get_links_from_page_js(url, selector, headless=True):
    opts = Options()
    if headless:
        opts.add_argument('--headless')
        opts.add_argument('--disable-gpu')
    opts.add_argument('--no-sandbox')
    opts.add_argument('--disable-dev-shm-usage')
    driver = webdriver.Chrome(options=opts)
    try:
        driver.get(url)
        time.sleep(2)
        items = driver.find_elements(By.CSS_SELECTOR, selector)
        results = []
        for item in items:
            text = item.text.strip()
            href = item.get_attribute('href') or item.get_attribute('data-href')
            results.append((text, href))
        return results
    except Exception as e:
        print(f"Lỗi get_links_from_page_js: {e}")
        return []
    finally:
        driver.quit()

def test_ignore_selectors(url, ignore_selectors, ignore_type, log_widget=None):
    opts = Options()
    opts.add_argument('--headless')
    opts.add_argument('--disable-gpu')
    driver = webdriver.Chrome(options=opts)
    try:
        driver.get(url)
        time.sleep(2)
        
        before_count = len(driver.find_elements(By.CSS_SELECTOR, "*"))
        
        if ignore_type == 'css':
            for selector in ignore_selectors:
                try:
                    elements = driver.find_elements(By.CSS_SELECTOR, selector)
                    for elem in elements:
                        driver.execute_script("arguments[0].remove();", elem)
                    log_to_gui(log_widget, f"Đã xóa {len(elements)} phần tử với selector: {selector}")
                except Exception as e:
                    log_to_gui(log_widget, f"Lỗi khi xóa {selector}: {e}")
        else:
            for js in ignore_selectors:
                try:
                    driver.execute_script(js)
                    log_to_gui(log_widget, f"Đã thực thi JS: {js[:50]}...")
                except Exception as e:
                    log_to_gui(log_widget, f"Lỗi khi thực thi JS: {e}")
        
        after_count = len(driver.find_elements(By.CSS_SELECTOR, "*"))
        log_to_gui(log_widget, f"Số phần tử trước: {before_count}, sau: {after_count}, đã xóa: {before_count - after_count}")
        
        return True
    except Exception as e:
        log_to_gui(log_widget, f"Lỗi khi test: {e}")
        return False
    finally:
        driver.quit()

def extract_main_blocks(driver, content_selectors):
    js_traverse = f'''
    function cleanText(text) {{
        return text.replace(/\\s+/g, ' ').trim();
    }}
    
    function collect(root){{
        var out = [];
        function walk(node){{
            if(!node) return;
            if(node.nodeType === Node.TEXT_NODE){{
                var t = cleanText(node.textContent);
                if(t && t.length > 0) out.push({{t:'text', v:t}});
            }} else if(node.nodeType === Node.ELEMENT_NODE){{
                if(node.tagName === 'IMG'){{
                    var src = node.src || node.getAttribute('data-src') || node.getAttribute('data-original');
                    if(src && src.startsWith('http')) out.push({{t:'img', v:src}});
                    return;
                }}
                if(node.tagName === 'BR'){{
                    out.push({{t:'text', v:'\\n'}});
                    return;
                }}
                if(node.tagName === 'H1' || node.tagName === 'H2' || node.tagName === 'H3'){{
                    var t = cleanText(node.textContent);
                    if(t) out.push({{t:'heading', v:t, level: parseInt(node.tagName[1])}});
                    return;
                }}
                if(node.tagName === 'P'){{
                    var t = cleanText(node.textContent);
                    if(t) out.push({{t:'paragraph', v:t}});
                    return;
                }}
                var children = node.childNodes;
                for(var i=0;i<children.length;i++) walk(children[i]);
            }}
        }}
        walk(root);
        return out;
    }}
    var root = null;
    var sels = {content_selectors};
    for(var i=0;i<sels.length;i++){{
        try{{ root = document.querySelector(sels[i]); if(root) break;}}catch(e){{}}
    }}
    if(!root) root = document.body;
    return collect(root);
    '''
    try:
        blocks = driver.execute_script(js_traverse)
        return blocks
    except Exception as e:
        print(f"Lỗi extract_main_blocks: {e}")
        return []

def save_image_from_url(img_url, absolute_url):
    try:
        if not img_url.startswith('http'):
            img_url = urljoin(absolute_url, img_url)
        
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
        resp = requests.get(img_url, headers=headers, stream=True, timeout=15)
        if resp.status_code == 200:
            content = resp.content
            content_type = resp.headers.get('content-type', '')
            
            if 'webp' in content_type or img_url.lower().endswith('.webp'):
                img = Image.open(BytesIO(content))
                fname = f"img_{int(time.time()*1000)}.png"
                fpath = os.path.join(IMAGE_TMP_DIR, fname)
                img.save(fpath, 'PNG')
                return fpath
            else:
                ext = 'jpg'
                if 'png' in content_type or img_url.lower().endswith('.png'):
                    ext = 'png'
                fname = f"img_{int(time.time()*1000)}.{ext}"
                fpath = os.path.join(IMAGE_TMP_DIR, fname)
                with open(fpath, 'wb') as f:
                    f.write(content)
                return fpath
    except Exception as e:
        print(f"Lỗi tải ảnh {img_url}: {e}")
    return None

def load_completed_links():
    if os.path.exists(COMPLETED_LINKS_FILE):
        try:
            with open(COMPLETED_LINKS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            return []
    return []

def save_completed_links(links):
    with open(COMPLETED_LINKS_FILE, 'w', encoding='utf-8') as f:
        json.dump(links, f, ensure_ascii=False, indent=2)

def clear_completed_links():
    if os.path.exists(COMPLETED_LINKS_FILE):
        os.remove(COMPLETED_LINKS_FILE)

def saved_links_path_for(base_url):
    parsed = urlparse(base_url)
    host = parsed.netloc.replace(':', '_') or 'default'
    return os.path.join(SAVED_LINKS_DIR, f"{host}.json")

def load_saved_links(base_url):
    path = saved_links_path_for(base_url)
    if os.path.exists(path):
        try:
            with open(path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            return []
    return []

def save_saved_links(base_url, links):
    path = saved_links_path_for(base_url)
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(links, f, ensure_ascii=False, indent=2)

def create_word_document_ext(config, output_path, log_widget=None, mode="append", headless=True):
    base_url = config['base_url']
    menu_selector = config['menu_selector']
    menu_selector_type = config['menu_selector_type']
    ignore_selectors = config['ignore_selectors']
    ignore_selectors_type = config['ignore_selectors_type']
    link_type = config['link_type']
    relative_base_url = config['relative_base_url']
    content_selectors = config.get('content_selectors', default_config()['content_selectors'])
    
    log_to_gui(log_widget, f"Bắt đầu lấy menu từ {base_url}")
    
    if menu_selector_type == 'css':
        links = get_links_from_page_css(base_url, menu_selector)
    else:
        links = get_links_from_page_js(base_url, menu_selector, headless)
    
    log_to_gui(log_widget, f"Tìm thấy {len(links)} link trong menu")
    
    if mode == "new":
        if os.path.exists(output_path):
            os.remove(output_path)
        clear_completed_links()
        doc = Document()
        doc.save(output_path)
        log_to_gui(log_widget, f"Tạo file mới: {output_path}")
    else:
        if os.path.exists(output_path):
            doc = Document(output_path)
        else:
            doc = Document()
            doc.save(output_path)
    
    completed_links = load_completed_links()
    saved_links = load_saved_links(base_url)
    
    existing_headings = [p.text.strip() for p in doc.paragraphs if p.style.name == 'Heading 2']
    
    new_links = []
    for text, href in links:
        if not href:
            continue
        if text in existing_headings or text in completed_links:
            log_to_gui(log_widget, f"Bỏ qua '{text}' đã tồn tại")
            continue
        new_links.append((text, href))
    
    log_to_gui(log_widget, f"Cần xử lý {len(new_links)} link mới")
    
    for idx, (text, href) in enumerate(new_links):
        if link_type == 'relative' and not href.startswith('http'):
            absolute = urljoin(relative_base_url or base_url, href)
        else:
            absolute = href
        
        log_to_gui(log_widget, f"[{idx+1}/{len(new_links)}] Đang xử lý: {text}")
        
        opts = Options()
        if headless:
            opts.add_argument('--headless')
            opts.add_argument('--disable-gpu')
        opts.add_argument('--no-sandbox')
        opts.add_argument('--disable-dev-shm-usage')
        
        driver = None
        try:
            driver = webdriver.Chrome(options=opts)
            driver.get(absolute)
            time.sleep(2)
            
            if ignore_selectors_type == 'css':
                for selector in ignore_selectors:
                    try:
                        elements = driver.find_elements(By.CSS_SELECTOR, selector)
                        for elem in elements:
                            driver.execute_script("arguments[0].remove();", elem)
                    except:
                        pass
            else:
                for js in ignore_selectors:
                    try:
                        driver.execute_script(js)
                    except:
                        pass
            
            blocks = extract_main_blocks(driver, content_selectors)
            
            if blocks:
                doc.add_heading(text, level=2)
                inserted_images = []
                
                for block in blocks:
                    if block.get('t') in ['text', 'paragraph']:
                        para_text = block.get('v', '')
                        if para_text:
                            p = doc.add_paragraph(para_text)
                            p.paragraph_format.space_after = Pt(6)
                    elif block.get('t') == 'heading':
                        level = block.get('level', 2)
                        if level == 1:
                            doc.add_heading(block.get('v', ''), level=1)
                        else:
                            doc.add_heading(block.get('v', ''), level=level)
                    elif block.get('t') == 'img':
                        img_path = save_image_from_url(block.get('v', ''), absolute)
                        if img_path and os.path.exists(img_path):
                            try:
                                doc.add_picture(img_path, width=Inches(6))
                                inserted_images.append(img_path)
                            except:
                                pass
                
                doc.save(output_path)
                log_to_gui(log_widget, f"✓ Đã thêm '{text}' vào tài liệu")
                
                completed_links.append(text)
                save_completed_links(completed_links)
                
                for img_path in inserted_images:
                    try:
                        os.remove(img_path)
                    except:
                        pass
            else:
                log_to_gui(log_widget, f"✗ Không lấy được nội dung cho {text}")
                
        except Exception as e:
            log_to_gui(log_widget, f"✗ Lỗi khi xử lý {absolute}: {e}")
        finally:
            if driver:
                driver.quit()
        
        time.sleep(1)
    
    save_saved_links(base_url, completed_links)
    log_to_gui(log_widget, f"Hoàn thành! Đã xử lý {len(completed_links)} link")

def show_ignore_tester(config, log_widget):
    test_window = Toplevel()
    test_window.title("Test bỏ qua phần tử")
    test_window.geometry("800x600")
    
    Label(test_window, text="URL test:").pack(pady=5)
    test_url_var = StringVar(value=config.get('base_url', ''))
    Entry(test_window, textvariable=test_url_var, width=80).pack(pady=5)
    
    Label(test_window, text="Selectors/JS (mỗi dòng một selector):").pack(pady=5)
    test_ignore_text = ScrolledText(test_window, width=90, height=10)
    test_ignore_text.pack(pady=5)
    test_ignore_text.insert('1.0', '\n'.join(config.get('ignore_selectors', [])))
    
    test_log = ScrolledText(test_window, width=90, height=20, state='disabled')
    test_log.pack(pady=5)
    
    def run_test():
        url = test_url_var.get().strip()
        ignore_list = [s.strip() for s in test_ignore_text.get('1.0', END).splitlines() if s.strip()]
        ignore_type = config.get('ignore_selectors_type', 'css')
        
        test_log.config(state='normal')
        test_log.delete('1.0', END)
        test_log.config(state='disabled')
        
        def test_log_func(msg):
            test_log.config(state='normal')
            test_log.insert(END, msg + "\n")
            test_log.see(END)
            test_log.config(state='disabled')
        
        test_ignore_selectors(url, ignore_list, ignore_type, test_log_func)
    
    Button(test_window, text="Test", command=run_test, bg='#4CAF50', fg='white').pack(pady=10)

def show_saved_websites():
    saved_window = Toplevel()
    saved_window.title("Các trang web đã lưu")
    saved_window.geometry("500x400")
    
    listbox = Listbox(saved_window, width=60, height=20)
    listbox.pack(pady=10, padx=10, fill='both', expand=True)
    
    config_files = [f.replace('.json', '') for f in os.listdir(CONFIG_DIR) if f.endswith('.json')]
    for cfg in config_files:
        listbox.insert(END, cfg)
    
    def load_selected():
        selection = listbox.curselection()
        if selection:
            name = listbox.get(selection[0])
            saved_window.destroy()
            return name
        return None
    
    Button(saved_window, text="Chọn", command=lambda: load_selected()).pack(pady=5)

def run_gui():
    config = load_config()
    root = Tk()
    root.title("Web to Word Automation - Tool thu thập nội dung website")
    root.geometry("1000x750")
    
    # Variables
    config_name_var = StringVar(value=config.get('config_name', ''))
    url_var = StringVar(value=config.get('base_url', ''))
    menu_selector_var = StringVar(value=config.get('menu_selector', 'ul#ul-search a'))
    menu_selector_type_var = StringVar(value=config.get('menu_selector_type', 'css'))
    ignore_selectors_type_var = StringVar(value=config.get('ignore_selectors_type', 'css'))
    output_var = StringVar(value=config.get('output_docx', 'output.docx'))
    link_type_var = StringVar(value=config.get('link_type', 'absolute'))
    relative_base_var = StringVar(value=config.get('relative_base_url', ''))
    
    # Notebook style layout
    main_frame = Frame(root)
    main_frame.pack(fill='both', expand=True, padx=10, pady=10)
    
    # Row 0: Config name
    Label(main_frame, text="Tên cấu hình:").grid(row=0, column=0, sticky='e', pady=5)
    Entry(main_frame, textvariable=config_name_var, width=30).grid(row=0, column=1, sticky='w', pady=5)
    
    Label(main_frame, text="Cấu hình đã lưu:").grid(row=0, column=2, sticky='e', pady=5)
    config_listbox = Listbox(main_frame, width=25, height=6)
    config_listbox.grid(row=0, column=3, rowspan=2, sticky='w', padx=5)
    
    def refresh_config_list():
        config_listbox.delete(0, END)
        for f in sorted([os.path.splitext(x)[0] for x in os.listdir(CONFIG_DIR) if x.endswith('.json')]):
            config_listbox.insert(END, f)
    refresh_config_list()
    
    # Row 1: URL
    Label(main_frame, text="URL trang chủ:").grid(row=1, column=0, sticky='e', pady=5)
    Entry(main_frame, textvariable=url_var, width=60).grid(row=1, column=1, columnspan=2, sticky='w', pady=5)
    
    # Row 2: Menu selector
    Label(main_frame, text="Menu selector:").grid(row=2, column=0, sticky='e', pady=5)
    Entry(main_frame, textvariable=menu_selector_var, width=60).grid(row=2, column=1, columnspan=2, sticky='w', pady=5)
    frame_ms = Frame(main_frame)
    frame_ms.grid(row=2, column=3, sticky='w')
    Radiobutton(frame_ms, text='CSS', variable=menu_selector_type_var, value='css').pack(side='left')
    Radiobutton(frame_ms, text='JavaScript', variable=menu_selector_type_var, value='javascript').pack(side='left')
    
    # Row 3: Content selectors
    Label(main_frame, text="Selector nội dung (mỗi dòng 1):").grid(row=3, column=0, sticky='ne', pady=5)
    content_text = ScrolledText(main_frame, width=60, height=4)
    content_text.grid(row=3, column=1, columnspan=2, sticky='w', pady=5)
    content_text.insert('1.0', '\n'.join(config.get('content_selectors', default_config()['content_selectors'])))
    
    # Row 4: Ignore selectors
    Label(main_frame, text="Bỏ qua (mỗi dòng 1):").grid(row=4, column=0, sticky='ne', pady=5)
    ignore_text = ScrolledText(main_frame, width=60, height=4)
    ignore_text.grid(row=4, column=1, columnspan=2, sticky='w', pady=5)
    ignore_text.insert('1.0', '\n'.join(config.get('ignore_selectors', [])))
    frame_ig = Frame(main_frame)
    frame_ig.grid(row=4, column=3, sticky='w')
    Radiobutton(frame_ig, text='CSS', variable=ignore_selectors_type_var, value='css').pack(side='left')
    Radiobutton(frame_ig, text='JavaScript', variable=ignore_selectors_type_var, value='javascript').pack(side='left')
    
    # Row 5: Output file
    Label(main_frame, text="File Word đích:").grid(row=5, column=0, sticky='e', pady=5)
    Entry(main_frame, textvariable=output_var, width=50).grid(row=5, column=1, sticky='w', pady=5)
    
    def choose_output_file():
        path = filedialog.asksaveasfilename(defaultextension='.docx', filetypes=[('Word', '*.docx')])
        if path:
            output_var.set(path)
    Button(main_frame, text='Chọn file', command=choose_output_file).grid(row=5, column=2, sticky='w', pady=5)
    
    # Row 6: Link type
    Label(main_frame, text="Loại link:").grid(row=6, column=0, sticky='e', pady=5)
    frame_link = Frame(main_frame)
    frame_link.grid(row=6, column=1, sticky='w', pady=5)
    Radiobutton(frame_link, text='Tuyệt đối', variable=link_type_var, value='absolute').pack(side='left')
    Radiobutton(frame_link, text='Tương đối', variable=link_type_var, value='relative').pack(side='left')
    Label(main_frame, text="Base URL (nếu tương đối):").grid(row=6, column=2, sticky='e', pady=5)
    Entry(main_frame, textvariable=relative_base_var, width=30).grid(row=6, column=3, sticky='w', pady=5)
    
    # Row 7: Log
    Label(main_frame, text="Log:").grid(row=7, column=0, sticky='ne', pady=5)
    log_text = ScrolledText(main_frame, width=90, height=12, state='disabled')
    log_text.grid(row=7, column=1, columnspan=3, sticky='w', pady=5)
    
    def update_config_from_gui(cfg):
        cfg['config_name'] = config_name_var.get().strip() or 'config'
        cfg['base_url'] = url_var.get().strip()
        cfg['menu_selector'] = menu_selector_var.get().strip()
        cfg['menu_selector_type'] = menu_selector_type_var.get()
        cfg['ignore_selectors'] = [s.strip() for s in ignore_text.get('1.0', END).splitlines() if s.strip()]
        cfg['ignore_selectors_type'] = ignore_selectors_type_var.get()
        cfg['output_docx'] = output_var.get().strip()
        cfg['link_type'] = link_type_var.get()
        cfg['relative_base_url'] = relative_base_var.get().strip()
        cfg['content_selectors'] = [s.strip() for s in content_text.get('1.0', END).splitlines() if s.strip()]
        if not cfg['content_selectors']:
            cfg['content_selectors'] = default_config()['content_selectors']
    
    def on_select_config(evt=None):
        sel = config_listbox.curselection()
        if sel:
            name = config_listbox.get(sel[0])
            loaded = load_config(name)
            config_name_var.set(loaded.get('config_name', name))
            url_var.set(loaded.get('base_url', ''))
            menu_selector_var.set(loaded.get('menu_selector', 'ul#ul-search a'))
            menu_selector_type_var.set(loaded.get('menu_selector_type', 'css'))
            ignore_text.delete('1.0', END)
            ignore_text.insert('1.0', '\n'.join(loaded.get('ignore_selectors', [])))
            ignore_selectors_type_var.set(loaded.get('ignore_selectors_type', 'css'))
            output_var.set(loaded.get('output_docx', 'output.docx'))
            link_type_var.set(loaded.get('link_type', 'absolute'))
            relative_base_var.set(loaded.get('relative_base_url', ''))
            content_text.delete('1.0', END)
            content_text.insert('1.0', '\n'.join(loaded.get('content_selectors', default_config()['content_selectors'])))
            log_to_gui(log_text, f"Đã nạp cấu hình: {name}")
    
    config_listbox.bind('<<ListboxSelect>>', on_select_config)
    
    def on_save():
        cfg = load_config()
        update_config_from_gui(cfg)
        save_config(cfg, cfg['config_name'])
        refresh_config_list()
        messagebox.showinfo('Thành công', f"Đã lưu cấu hình '{cfg['config_name']}'")
    
    def on_load():
        sel = config_listbox.curselection()
        if sel:
            name = config_listbox.get(sel[0])
            on_select_config()
        else:
            messagebox.showwarning('Chú ý', 'Vui lòng chọn cấu hình trong danh sách!')
    
    def on_test_menu():
        cfg = load_config()
        update_config_from_gui(cfg)
        log_text.config(state='normal')
        log_text.delete('1.0', END)
        log_text.config(state='disabled')
        
        try:
            if cfg['menu_selector_type'] == 'css':
                links = get_links_from_page_css(cfg['base_url'], cfg['menu_selector'])
            else:
                links = get_links_from_page_js(cfg['base_url'], cfg['menu_selector'], headless=True)
            
            log_to_gui(log_text, f"Tìm thấy {len(links)} link trong menu:")
            for i, (t, h) in enumerate(links[:20]):
                log_to_gui(log_text, f"  {i+1}. {t[:50]} -> {h}")
            if len(links) > 20:
                log_to_gui(log_text, f"  ... và {len(links)-20} link khác")
        except Exception as e:
            log_to_gui(log_text, f"Lỗi: {e}")
    
    def on_test_ignore():
        cfg = load_config()
        update_config_from_gui(cfg)
        show_ignore_tester(cfg, log_text)
    
    def on_test_download():
        cfg = load_config()
        update_config_from_gui(cfg)
        
        test_window = Toplevel(root)
        test_window.title("Test tải một link")
        test_window.geometry("600x400")
        
        Label(test_window, text="URL cần test:").pack(pady=5)
        test_url = Entry(test_window, width=70)
        test_url.pack(pady=5)
        
        Label(test_window, text="Tiêu đề:").pack(pady=5)
        test_title = Entry(test_window, width=70)
        test_title.pack(pady=5)
        
        test_log = ScrolledText(test_window, width=80, height=15, state='disabled')
        test_log.pack(pady=5)
        
        def run_test():
            url = test_url.get().strip()
            title = test_title.get().strip() or "Test Page"
            
            if not url:
                messagebox.showwarning('Chú ý', 'Vui lòng nhập URL!')
                return
            
            def test_log_func(msg):
                test_log.config(state='normal')
                test_log.insert(END, msg + "\n")
                test_log.see(END)
                test_log.config(state='disabled')
            
            test_log_func(f"Đang test: {title} - {url}")
            
            opts = Options()
            opts.add_argument('--headless')
            opts.add_argument('--disable-gpu')
            driver = webdriver.Chrome(options=opts)
            
            try:
                driver.get(url)
                time.sleep(2)
                
                for selector in cfg['ignore_selectors']:
                    try:
                        elements = driver.find_elements(By.CSS_SELECTOR, selector)
                        for elem in elements:
                            driver.execute_script("arguments[0].remove();", elem)
                        if elements:
                            test_log_func(f"Đã xóa {len(elements)} phần tử với selector: {selector}")
                    except:
                        pass
                
                blocks = extract_main_blocks(driver, cfg.get('content_selectors', default_config()['content_selectors']))
                
                if blocks:
                    test_log_func(f"Lấy được {len(blocks)} block nội dung")
                    text_blocks = sum(1 for b in blocks if b.get('t') in ['text', 'paragraph'])
                    img_blocks = sum(1 for b in blocks if b.get('t') == 'img')
                    test_log_func(f"  - Text blocks: {text_blocks}")
                    test_log_func(f"  - Image blocks: {img_blocks}")
                    
                    test_doc = Document()
                    test_doc.add_heading(title, level=2)
                    
                    for block in blocks[:10]:
                        if block.get('t') in ['text', 'paragraph']:
                            test_doc.add_paragraph(block.get('v', '')[:200])
                    
                    test_path = os.path.join(DATA_DIR, "test_output.docx")
                    test_doc.save(test_path)
                    test_log_func(f"Đã lưu file test tại: {test_path}")
                else:
                    test_log_func("Không lấy được nội dung!")
                    
            except Exception as e:
                test_log_func(f"Lỗi: {e}")
            finally:
                driver.quit()
        
        Button(test_window, text="Chạy test", command=run_test, bg='#4CAF50', fg='white').pack(pady=10)
    
    def on_start_new():
        cfg = load_config()
        update_config_from_gui(cfg)
        create_word_document_ext(cfg, cfg['output_docx'], log_widget=log_text, mode="new", headless=True)
        messagebox.showinfo('Hoàn thành', 'Đã tạo file Word mới với toàn bộ nội dung!')
    
    def on_start_append():
        cfg = load_config()
        update_config_from_gui(cfg)
        create_word_document_ext(cfg, cfg['output_docx'], log_widget=log_text, mode="append", headless=True)
        messagebox.showinfo('Hoàn thành', 'Đã thêm nội dung mới vào file Word!')
    
    def on_forget_links():
        cfg = load_config()
        update_config_from_gui(cfg)
        if messagebox.askyesno('Xác nhận', 'Bạn có chắc muốn xóa danh sách link đã tải?\nCác link này sẽ được tải lại từ đầu!'):
            clear_completed_links()
            log_to_gui(log_text, "Đã xóa danh sách link đã tải!")
            messagebox.showinfo('Thành công', 'Đã xóa danh sách link!')
    
    def on_clear_log():
        log_text.config(state='normal')
        log_text.delete('1.0', END)
        log_text.config(state='disabled')
    
    # Buttons
    button_frame = Frame(main_frame)
    button_frame.grid(row=8, column=0, columnspan=4, pady=10)
    
    Button(button_frame, text='💾 Lưu cấu hình', command=on_save, width=14, bg='#2196F3', fg='white').pack(side='left', padx=5)
    Button(button_frame, text='📂 Tải cấu hình', command=on_load, width=14, bg='#FF9800', fg='white').pack(side='left', padx=5)
    Button(button_frame, text='🔍 Test Menu', command=on_test_menu, width=14, bg='#9C27B0', fg='white').pack(side='left', padx=5)
    Button(button_frame, text='🧪 Test bỏ qua', command=on_test_ignore, width=14, bg='#E91E63', fg='white').pack(side='left', padx=5)
    Button(button_frame, text='📝 Test tải 1 link', command=on_test_download, width=14, bg='#3F51B5', fg='white').pack(side='left', padx=5)
    
    button_frame2 = Frame(main_frame)
    button_frame2.grid(row=9, column=0, columnspan=4, pady=5)
    
    Button(button_frame2, text='🆕 Tạo mới (tải lại toàn bộ)', command=on_start_new, width=20, bg='#4CAF50', fg='white').pack(side='left', padx=5)
    Button(button_frame2, text='➕ Thêm mới (chỉ tải link chưa có)', command=on_start_append, width=25, bg='#8BC34A', fg='white').pack(side='left', padx=5)
    Button(button_frame2, text='🗑️ Xóa link đã tải', command=on_forget_links, width=16, bg='#F44336', fg='white').pack(side='left', padx=5)
    Button(button_frame2, text='🗑️ Xóa log', command=on_clear_log, width=12, bg='#9E9E9E', fg='white').pack(side='left', padx=5)
    
    root.mainloop()

if __name__ == '__main__':
    run_gui()