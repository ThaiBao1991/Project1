import os
import time
import json
import requests
import sys
from urllib.parse import urljoin, urlparse
import re
from bs4 import BeautifulSoup
from tkinter import Tk, filedialog, messagebox, StringVar, Entry, Button, Label, Listbox, END, Frame
from tkinter.scrolledtext import ScrolledText
from tkinter import Radiobutton
from docx import Document
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from pathlib import Path
from PIL import Image
from io import BytesIO

CONFIG_FILE = "config.json"
CONFIG_DIR = "configs"
DATA_DIR = os.path.join(os.path.dirname(__file__), "data")
SAVED_LINKS_DIR = os.path.join(DATA_DIR, "saved_links")
IMAGE_TMP_DIR = os.path.join(DATA_DIR, "images")

def default_config():
    return {
        "config_name": "Cau hinh mac dinh",
        "base_url": "https://vi.extendoffice.com",
        "menu_selector": "ul#ul-search a",
        "menu_selector_type": "css",  # css | javascript
        "ignore_selectors": [
            ".uk-margin-remove-last-child.custom",
            "div.uk-margin-remove-last-child.custom style"
        ],
        "ignore_selectors_type": "css",  # css | javascript
        "output_docx": "output.docx",
        "link_type": "absolute",
        "relative_base_url": ""
    }

if not os.path.exists(CONFIG_DIR):
    os.makedirs(CONFIG_DIR)

def save_config(config, name=None):
    if name:
        path = os.path.join(CONFIG_DIR, f"{name}.json")
    else:
        path = CONFIG_FILE
    with open(path, "w", encoding="utf-8") as f:
        json.dump(config, f, ensure_ascii=False, indent=2)

def load_config(name=None):
    if name:
        path = os.path.join(CONFIG_DIR, f"{name}.json")
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return default_config()

def log_to_gui(log_widget, msg):
    if not log_widget:
        print(msg)
        return
    log_widget.config(state="normal")
    log_widget.insert(END, msg + "\n")
    log_widget.see(END)
    log_widget.config(state="disabled")

def get_links_from_page_css(url, selector):
    resp = requests.get(url)
    soup = BeautifulSoup(resp.content, 'html.parser')
    items = soup.select(selector)
    results = []
    for i in items:
        text = i.get_text(strip=True)
        href = i.get('href') or i.get('data-href') or i.attrs.get('data-href')
        if not href:
            onclick = i.get('onclick') or i.attrs.get('onclick')
            if onclick:
                m = re.search(r"location\.href\s*=\s*['\"]([^'\"]+)['\"]", onclick)
                if m:
                    href = m.group(1)
        if href and href.strip().startswith('javascript:'):
            href = None
        results.append((text, href))
    return results

def get_links_from_page_js(url, selector, headless=True):
    opts = Options()
    if headless:
        opts.add_argument('--headless')
        opts.add_argument('--disable-gpu')
    driver = webdriver.Chrome(options=opts)
    driver.get(url)
    # Use selector as CSS selector in JS mode for convenience
    script = f"""
    var res = [];
    var els = document.querySelectorAll('{selector}');
    for(var i=0;i<els.length;i++){{
      var el = els[i];
      res.push({{text: el.innerText||'', href: el.href||el.getAttribute('href')}});
    }}
    return res;
    """
    try:
        items = driver.execute_script(script)
        return [(it.get('text','').strip(), it.get('href')) for it in items]
    finally:
        driver.quit()

def apply_ignore_rules(driver, ignore_list, ignore_type):
    if not ignore_list:
        return
    if ignore_type == 'css':
        for s in ignore_list:
            safe = s.replace("'","\\'")
            driver.execute_script(f"document.querySelectorAll('{safe}').forEach(e=>e.remove())")
    else:
        # treat each string as JS snippet to run
        for js in ignore_list:
            try:
                driver.execute_script(js)
            except Exception:
                continue

def ensure_saved_links_dir():
    if not os.path.exists(DATA_DIR):
        os.makedirs(DATA_DIR)
    if not os.path.exists(SAVED_LINKS_DIR):
        os.makedirs(SAVED_LINKS_DIR)
    if not os.path.exists(IMAGE_TMP_DIR):
        os.makedirs(IMAGE_TMP_DIR)

def saved_links_path_for(base_url):
    parsed = urlparse(base_url)
    host = parsed.netloc.replace(':','_') or 'default'
    ensure_saved_links_dir()
    return os.path.join(SAVED_LINKS_DIR, f"{host}.json")

def load_saved_links(base_url):
    path = saved_links_path_for(base_url)
    if os.path.exists(path):
        try:
            with open(path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            return []
    return []

def save_saved_links(base_url, links):
    path = saved_links_path_for(base_url)
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(links, f, ensure_ascii=False, indent=2)

# ensure data dirs exist on import
ensure_saved_links_dir()

def extract_main_blocks(driver):
        # Return ordered blocks of text/imgs from main/article or body
        candidates = [
                "article",
                "main",
                "div[class*='content']",
                "div[class*='article']",
                "div[class*='post']",
        ]
        js_traverse = '''
        function collect(root){
            var out = [];
            function walk(node){
                if(!node) return;
                if(node.nodeType === Node.TEXT_NODE){
                    var t = node.textContent.replace(/\s+/g,' ').trim();
                    if(t) out.push({t:'text', v:t});
                } else if(node.nodeType === Node.ELEMENT_NODE){
                    if(node.tagName === 'IMG'){
                        out.push({t:'img', v: node.src || node.getAttribute('data-src') || ''});
                        return;
                    }
                    if(node.tagName === 'BR'){
                        out.push({t:'text', v:'\n'});
                        return;
                    }
                    var children = node.childNodes;
                    for(var i=0;i<children.length;i++) walk(children[i]);
                }
            }
            walk(root);
            return out;
        }
        var root = null;
        var sels = %s;
        for(var i=0;i<sels.length;i++){
            try{ root = document.querySelector(sels[i]); if(root) break;}catch(e){}
        }
        if(!root) root = document.body;
        return collect(root);
        ''' % (str(candidates))
        try:
                blocks = driver.execute_script(js_traverse)
                return blocks
        except Exception:
                return []

def create_word_document_ext(base_url, menu_selector, menu_selector_type, ignore_selectors, ignore_selectors_type, document, output_path, link_type="absolute", relative_base_url="", log_widget=None, headless=True):
    # Get list of links depending on selector type
    log_to_gui(log_widget, f"Lấy menu từ {base_url} using {menu_selector_type}")
    try:
        if menu_selector_type == 'css':
            links = get_links_from_page_css(base_url, menu_selector)
        else:
            links = get_links_from_page_js(base_url, menu_selector, headless=headless)
    except Exception as e:
        log_to_gui(log_widget, f"Lỗi khi lấy menu: {e}")
        links = []

    log_to_gui(log_widget, f"Tìm thấy {len(links)} link trong menu")
    existing = [p.text.strip() for p in document.paragraphs if p.style.name == 'Heading 2' and p.text]
    saved_links = load_saved_links(base_url)
    saved_texts = [s.get('text') for s in saved_links]
    new_saved = list(saved_links)

    for i, (text, href) in enumerate(links):
        if not href:
            continue
        if text in existing or text in saved_texts:
            log_to_gui(log_widget, f"Bỏ qua '{text}' đã tồn tại")
            continue
        if link_type == 'relative' and not href.startswith('http'):
            absolute = urljoin(relative_base_url or base_url, href)
        else:
            absolute = href
        log_to_gui(log_widget, f"Xử lý {i}: {text} -> {absolute}")

        # load page with selenium to honor JS and ignore rules
        opts = Options()
        if headless:
            opts.add_argument('--headless')
            opts.add_argument('--disable-gpu')
        driver = webdriver.Chrome(options=opts)
        try:
            driver.get(absolute)
            apply_ignore_rules(driver, ignore_selectors, ignore_selectors_type)
            time.sleep(0.5)
            blocks = extract_main_blocks(driver)
            if blocks:
                document.add_heading(text, level=2)
                inserted_images = []
                for b in blocks:
                    if b.get('t') == 'text':
                        # preserve paragraphs by double newline
                        for para in b.get('v','').split('\n\n'):
                            if para.strip():
                                document.add_paragraph(para.strip())
                    elif b.get('t') == 'img':
                        img_url = b.get('v')
                        if not img_url:
                            continue
                        if not img_url.startswith('http'):
                            img_url = urljoin(absolute, img_url)
                        try:
                            resp = requests.get(img_url, stream=True, timeout=15)
                            if resp.status_code == 200:
                                content = resp.content
                                ct = resp.headers.get('content-type','')
                                ext = 'jpg'
                                if 'png' in ct or img_url.lower().endswith('.png'):
                                    ext = 'png'
                                elif 'webp' in ct or img_url.lower().endswith('.webp'):
                                    ext = 'png'
                                elif 'jpeg' in ct or img_url.lower().endswith(('.jpg','.jpeg')):
                                    ext = 'jpg'
                                fname = f"img_{int(time.time()*1000)}.{ext}"
                                fpath = os.path.join(IMAGE_TMP_DIR, fname)
                                if ext == 'png' and ('webp' in ct or img_url.lower().endswith('.webp')):
                                    try:
                                        im = Image.open(BytesIO(content)).convert('RGB')
                                        im.save(fpath, 'PNG')
                                    except Exception:
                                        continue
                                else:
                                    with open(fpath, 'wb') as f:
                                        f.write(content)
                                try:
                                    document.add_picture(fpath)
                                    inserted_images.append(fpath)
                                except Exception:
                                    pass
                        except Exception:
                            continue
                document.save(output_path)
                log_to_gui(log_widget, f"Đã thêm '{text}' vào tài liệu")
                if not any(s.get('text') == text for s in new_saved):
                    new_saved.append({'text': text, 'href': absolute})
                for p in inserted_images:
                    try:
                        os.remove(p)
                    except Exception:
                        pass
            else:
                log_to_gui(log_widget, f"Không có nội dung lấy được cho {absolute}")
        except Exception as e:
            log_to_gui(log_widget, f"Lỗi khi xử lý {absolute}: {e}")
        finally:
            driver.quit()
    # persist saved links
    try:
        save_saved_links(base_url, new_saved)
    except Exception:
        pass

# --- GUI ---
def run_gui():
    config = load_config()
    root = Tk()
    root.title("Web to Word Automation")

    # Variables
    config_name_var = StringVar(value=config.get('config_name',''))
    url_var = StringVar(value=config.get('base_url',''))
    menu_selector_var = StringVar(value=config.get('menu_selector','ul#ul-search a'))
    menu_selector_type_var = StringVar(value=config.get('menu_selector_type','css'))
    # ignore selectors stored as multi-line text (each line is a selector or JS snippet)
    ignore_selectors_type_var = StringVar(value=config.get('ignore_selectors_type','css'))
    output_var = StringVar(value=config.get('output_docx','output.docx'))
    link_type_var = StringVar(value=config.get('link_type','absolute'))
    relative_base_var = StringVar(value=config.get('relative_base_url',''))
    headless_var = StringVar(value='1')

    Label(root, text="Tên cấu hình:").grid(row=0, column=0, sticky='e')
    Entry(root, textvariable=config_name_var, width=30).grid(row=0, column=1, sticky='w')

    Label(root, text="Chọn cấu hình đã lưu:").grid(row=0, column=2, sticky='e')
    config_listbox = Listbox(root, width=25, height=4)
    config_listbox.grid(row=0, column=3, rowspan=2, sticky='w')

    def refresh_config_list():
        config_listbox.delete(0, END)
        for f in sorted([os.path.splitext(x)[0] for x in os.listdir(CONFIG_DIR) if x.endswith('.json')]):
            config_listbox.insert(END, f)
    refresh_config_list()

    def update_config_from_gui(cfg):
        cfg['config_name'] = config_name_var.get().strip() or 'config'
        cfg['base_url'] = url_var.get().strip()
        cfg['menu_selector'] = menu_selector_var.get().strip()
        cfg['menu_selector_type'] = menu_selector_type_var.get()
        raw = ignore_text.get('1.0', END)
        cfg['ignore_selectors'] = [s.strip() for s in raw.splitlines() if s.strip()]
        cfg['ignore_selectors_type'] = ignore_selectors_type_var.get()
        cfg['output_docx'] = output_var.get().strip()
        cfg['link_type'] = link_type_var.get()
        cfg['relative_base_url'] = relative_base_var.get().strip()

    def on_select_config(evt=None):
        sel = config_listbox.curselection()
        if sel:
            name = config_listbox.get(sel[0])
            loaded = load_config(name)
            config_name_var.set(loaded.get('config_name',name))
            url_var.set(loaded.get('base_url',''))
            menu_selector_var.set(loaded.get('menu_selector','ul#ul-search a'))
            menu_selector_type_var.set(loaded.get('menu_selector_type','css'))
            ignore_text.delete('1.0', END)
            ignore_text.insert('1.0', '\n'.join(loaded.get('ignore_selectors',[])))
            ignore_selectors_type_var.set(loaded.get('ignore_selectors_type','css'))
            output_var.set(loaded.get('output_docx','output.docx'))
            link_type_var.set(loaded.get('link_type','absolute'))
            relative_base_var.set(loaded.get('relative_base_url',''))

    config_listbox.bind('<<ListboxSelect>>', on_select_config)

    Label(root, text="Link trang chủ:").grid(row=1, column=0, sticky='e')
    Entry(root, textvariable=url_var, width=50).grid(row=1, column=1, columnspan=2, sticky='w')

    Label(root, text="Menu selector:").grid(row=2, column=0, sticky='e')
    Entry(root, textvariable=menu_selector_var, width=50).grid(row=2, column=1, columnspan=2, sticky='w')
    Frame_ms = Frame(root)
    Frame_ms.grid(row=2, column=3, sticky='w')
    Radiobutton(Frame_ms, text='CSS', variable=menu_selector_type_var, value='css').pack(side='left')
    Radiobutton(Frame_ms, text='JavaScript', variable=menu_selector_type_var, value='javascript').pack(side='left')

    Label(root, text="Yếu tố bỏ qua (mỗi dòng 1 selector/JS):").grid(row=3, column=0, sticky='ne')
    ignore_text = ScrolledText(root, width=50, height=4)
    ignore_text.grid(row=3, column=1, columnspan=2, sticky='w')
    ignore_text.insert('1.0', "\n".join(config.get('ignore_selectors',[])))
    Frame_ig = Frame(root)
    Frame_ig.grid(row=3, column=3, sticky='w')
    Radiobutton(Frame_ig, text='CSS', variable=ignore_selectors_type_var, value='css').pack(side='left')
    Radiobutton(Frame_ig, text='JavaScript', variable=ignore_selectors_type_var, value='javascript').pack(side='left')

    Label(root, text="Tên file Word:").grid(row=4, column=0, sticky='e')
    Entry(root, textvariable=output_var, width=40).grid(row=4, column=1, sticky='w')
    def choose_output_file():
        path = filedialog.asksaveasfilename(defaultextension='.docx', filetypes=[('Word', '*.docx')])
        if path:
            output_var.set(path)
    Button(root, text='Chọn file', command=choose_output_file).grid(row=4, column=2, sticky='w')

    Label(root, text='Loại link menu:').grid(row=5, column=0, sticky='e')
    Frame_link = Frame(root)
    Frame_link.grid(row=5, column=1, sticky='w')
    Radiobutton(Frame_link, text='Tuyệt đối', variable=link_type_var, value='absolute').pack(side='left')
    Radiobutton(Frame_link, text='Tương đối', variable=link_type_var, value='relative').pack(side='left')
    Label(root, text='Base URL (nếu tương đối):').grid(row=5, column=2, sticky='e')
    Entry(root, textvariable=relative_base_var, width=30).grid(row=5, column=3, sticky='w')

    Label(root, text='Log:').grid(row=6, column=0, sticky='ne')
    log_text = ScrolledText(root, width=90, height=14, state='disabled')
    log_text.grid(row=6, column=1, columnspan=3, sticky='w')

    def on_start():
        cfg = load_config()
        update_config_from_gui(cfg)
        out = cfg['output_docx']
        if os.path.exists(out):
            doc = Document(out)
        else:
            doc = Document()
            doc.add_paragraph('')
            doc.save(out)
        create_word_document_ext(cfg['base_url'], cfg['menu_selector'], cfg['menu_selector_type'], cfg['ignore_selectors'], cfg['ignore_selectors_type'], doc, out, cfg['link_type'], cfg['relative_base_url'], log_widget=log_text, headless=True)
        messagebox.showinfo('Xong', 'Đã hoàn thành copy dữ liệu!')

    def on_save():
        cfg = load_config()
        update_config_from_gui(cfg)
        save_config(cfg, cfg['config_name'])
        refresh_config_list()
        messagebox.showinfo('Lưu cấu hình', f"Đã lưu cấu hình vào {cfg['config_name']}.json!")

    def on_load():
        sel = config_listbox.curselection()
        if sel:
            name = config_listbox.get(sel[0])
            loaded = load_config(name)
            config_name_var.set(loaded.get('config_name', name))
            url_var.set(loaded.get('base_url',''))
            menu_selector_var.set(loaded.get('menu_selector','ul#ul-search a'))
            menu_selector_type_var.set(loaded.get('menu_selector_type','css'))
            ignore_text.delete('1.0', END)
            ignore_text.insert('1.0', '\n'.join(loaded.get('ignore_selectors',[])))
            ignore_selectors_type_var.set(loaded.get('ignore_selectors_type','css'))
            output_var.set(loaded.get('output_docx','output.docx'))
            link_type_var.set(loaded.get('link_type','absolute'))
            relative_base_var.set(loaded.get('relative_base_url',''))
            log_to_gui(log_text, f"Đã nạp cấu hình {name}")
        else:
            messagebox.showinfo('Nạp cấu hình', 'Chọn cấu hình trong danh sách!')

    def on_test_menu():
        cfg = load_config()
        update_config_from_gui(cfg)
        log_text.config(state='normal')
        log_text.delete(1.0, END)
        log_text.config(state='disabled')
        try:
            if cfg['menu_selector_type'] == 'css':
                links = get_links_from_page_css(cfg['base_url'], cfg['menu_selector'])
            else:
                links = get_links_from_page_js(cfg['base_url'], cfg['menu_selector'], headless=True)
            log_to_gui(log_text, f"Tìm thấy {len(links)} link trong menu:")
            for i,(t,h) in enumerate(links):
                log_to_gui(log_text, f"Link {i}: {t} - {h}")
        except Exception as e:
            log_to_gui(log_text, f"Lỗi khi lấy menu: {e}")

    def on_test_content():
        cfg = load_config()
        update_config_from_gui(cfg)
        try:
            if cfg['menu_selector_type'] == 'css':
                links = get_links_from_page_css(cfg['base_url'], cfg['menu_selector'])
            else:
                links = get_links_from_page_js(cfg['base_url'], cfg['menu_selector'], headless=True)
            if not links:
                log_to_gui(log_text, 'Không tìm thấy link để test.')
                return
            text, href = links[0]
            if cfg['link_type'] == 'relative' and not href.startswith('http'):
                absolute = urljoin(cfg['relative_base_url'] or cfg['base_url'], href)
            else:
                absolute = href
            log_to_gui(log_text, f"Test lấy nội dung: {absolute}")
            opts = Options()
            opts.add_argument('--headless')
            driver = webdriver.Chrome(options=opts)
            try:
                driver.get(absolute)
                apply_ignore_rules(driver, cfg['ignore_selectors'], cfg['ignore_selectors_type'])
                blocks = extract_main_blocks(driver)
                if blocks:
                    texts = sum(1 for b in blocks if b.get('t') == 'text')
                    imgs = sum(1 for b in blocks if b.get('t') == 'img')
                    total_chars = sum(len(b.get('v','')) for b in blocks if b.get('t') == 'text')
                    log_to_gui(log_text, f"Lấy được {texts} text-blocks, {imgs} ảnh, tổng {total_chars} chars")
                else:
                    log_to_gui(log_text, 'Không lấy được nội dung')
            finally:
                driver.quit()
        except Exception as e:
            log_to_gui(log_text, f"Lỗi khi test nội dung: {e}")

    def clear_log():
        log_text.config(state='normal')
        log_text.delete(1.0, END)
        log_text.config(state='disabled')

    def forget_saved_links():
        base = url_var.get().strip() or config.get('base_url','')
        path = saved_links_path_for(base)
        if os.path.exists(path):
            try:
                os.remove(path)
                log_to_gui(log_text, f"Đã xóa saved links cho {base}")
            except Exception as e:
                log_to_gui(log_text, f"Lỗi khi xóa saved links: {e}")
        else:
            log_to_gui(log_text, f"Không tìm thấy saved links cho {base}")

    def on_test_download_one():
        cfg = load_config()
        update_config_from_gui(cfg)
        try:
            if cfg['menu_selector_type'] == 'css':
                links = get_links_from_page_css(cfg['base_url'], cfg['menu_selector'])
            else:
                links = get_links_from_page_js(cfg['base_url'], cfg['menu_selector'], headless=True)
            if not links:
                log_to_gui(log_text, 'Không tìm thấy link để test.')
                return
            text, href = links[0]
            if cfg['link_type'] == 'relative' and not href.startswith('http'):
                absolute = urljoin(cfg['relative_base_url'] or cfg['base_url'], href)
            else:
                absolute = href
            log_to_gui(log_text, f"Test download link đầu tiên: {absolute}")
            opts = Options(); opts.add_argument('--headless')
            driver = webdriver.Chrome(options=opts)
            try:
                driver.get(absolute)
                apply_ignore_rules(driver, cfg['ignore_selectors'], cfg['ignore_selectors_type'])
                blocks = extract_main_blocks(driver)
                if not blocks:
                    log_to_gui(log_text, 'Không lấy được nội dung để test')
                    return
                # create temp doc
                ensure_saved_links_dir()
                tmp_doc = Document()
                tmp_doc.add_heading(text, level=2)
                inserted = []
                for b in blocks:
                    if b.get('t') == 'text':
                        for para in b.get('v','').split('\n\n'):
                            if para.strip():
                                tmp_doc.add_paragraph(para.strip())
                    elif b.get('t') == 'img':
                        img_url = b.get('v')
                        if not img_url:
                            continue
                        if not img_url.startswith('http'):
                            img_url = urljoin(absolute, img_url)
                        try:
                            resp = requests.get(img_url, stream=True, timeout=15)
                            if resp.status_code == 200:
                                content = resp.content
                                ct = resp.headers.get('content-type','')
                                ext = 'jpg'
                                if 'png' in ct or img_url.lower().endswith('.png'):
                                    ext = 'png'
                                elif 'webp' in ct or img_url.lower().endswith('.webp'):
                                    ext = 'png'
                                fname = f"img_test_{int(time.time()*1000)}.{ext}"
                                fpath = os.path.join(IMAGE_TMP_DIR, fname)
                                if ext == 'png' and ('webp' in ct or img_url.lower().endswith('.webp')):
                                    try:
                                        im = Image.open(BytesIO(content)).convert('RGB')
                                        im.save(fpath, 'PNG')
                                    except Exception:
                                        continue
                                else:
                                    with open(fpath, 'wb') as f:
                                        f.write(content)
                                try:
                                    tmp_doc.add_picture(fpath)
                                    inserted.append(fpath)
                                except Exception:
                                    pass
                        except Exception:
                            continue
                out = os.path.join(DATA_DIR, 'test_single.docx')
                tmp_doc.save(out)
                log_to_gui(log_text, f"Đã lưu file test vào {out}")
                try:
                    os.startfile(out)
                except Exception:
                    pass
                for p in inserted:
                    try:
                        os.remove(p)
                    except Exception:
                        pass
            finally:
                driver.quit()
        except Exception as e:
            log_to_gui(log_text, f"Lỗi khi test download: {e}")

    Button(root, text='Start', command=on_start, width=12, bg='#4CAF50', fg='white').grid(row=7, column=0, pady=10)
    Button(root, text='Save', command=on_save, width=12, bg='#2196F3', fg='white').grid(row=7, column=1, pady=10)
    Button(root, text='Load', command=on_load, width=12, bg='#FFC107', fg='black').grid(row=7, column=2, pady=10)
    Button(root, text='Test Menu', command=on_test_menu, width=12, bg='#9C27B0', fg='white').grid(row=7, column=3, pady=10)
    Button(root, text='Test Content', command=on_test_content, width=12, bg='#FF5722', fg='white').grid(row=7, column=4, pady=10)
    Button(root, text='Clear Log', command=clear_log, width=12, bg='#9E9E9E', fg='white').grid(row=8, column=0, pady=6)
    Button(root, text='Forget Saved Links', command=forget_saved_links, width=16, bg='#607D8B', fg='white').grid(row=8, column=1, pady=6)
    Button(root, text='Test Download 1', command=on_test_download_one, width=14, bg='#3E2723', fg='white').grid(row=8, column=2, pady=6)

    root.mainloop()

if __name__ == '__main__':
    run_gui()
