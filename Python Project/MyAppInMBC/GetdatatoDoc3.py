import os
import time
import pyautogui
import pywinauto
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait 
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
from urllib.parse import urljoin
from bs4 import BeautifulSoup
import requests
from tkinter import Tk, filedialog, messagebox
from pathlib import Path
from pywinauto.findwindows import ElementNotFoundError
from PIL import ImageGrab, ImageChops
import win32com.client
from pywinauto.application import Application
import pyperclip
from selenium.common.exceptions import TimeoutException

user_dir = Path("C:/Users/12953 bao/Desktop/desktop/work/Project/Python/BasicLearnPython/W3schools")
output_path = user_dir / "output.docx"

def select_image_file(default_path):
    root = Tk()
    root.withdraw()
    root.attributes('-topmost', True)
    file_path = filedialog.askopenfilename(initialdir=os.path.dirname(default_path), title="Select Image File",
                                           filetypes=(("PNG files", "*.png"), ("All files", "*.*")))
    root.destroy()
    return file_path

def select_word_file():
    root = Tk()
    root.withdraw()
    root.attributes('-topmost', True)
    file_path = filedialog.askopenfilename(title="Select Word File",
                                           filetypes=(("Word files", "*.docx"), ("All files", "*.*")))
    root.destroy()
    return file_path

# Hàm chụp ảnh màn hình của cửa sổ
def capture_screenshot(dlg):
    rect = dlg.rectangle()
    screenshot = ImageGrab.grab(bbox=(rect.left, rect.top, rect.right, rect.bottom))
    return screenshot

# Hàm kiểm tra sự thay đổi của cửa sổ
def has_window_changed(dlg, initial_screenshot):
    current_screenshot = capture_screenshot(dlg)
    diff = ImageChops.difference(initial_screenshot, current_screenshot)
    return diff.getbbox() is not None

def wait_for_window(app,title_re, timeout=10):
    start_time = time.time()
    while time.time() - start_time < timeout:
        try:
            dlg = app.window(title_re=title_re)
            if dlg.exists(timeout=1):
                return dlg
        except ElementNotFoundError:
            pass
        time.sleep(0.5)
    return None

def wait_for_save_completion(file_path, timeout=30):
    start_time = time.time()
    file_size = -1  # Kích thước file ban đầu

    while time.time() - start_time < timeout:
        try:
            if os.path.exists(file_path):
                current_size = os.path.getsize(file_path)
                if current_size == file_size: # Kiểm tra xem kích thước file có thay đổi nữa không
                    return True  # Kích thước file ổn định, coi như đã lưu xong
                else:
                    file_size = current_size
                    time.sleep(1) # Chờ 1s rồi kiểm tra lại
            else:
                print("File không tồn tại.")
                return False
        except Exception as e:
            print(f"Lỗi trong quá trình kiểm tra: {e}")
            return False
    return False

def close_word_document(dlg):
    try:
        dlg.close()
        print("Đã đóng cửa sổ Word.")
    except pywinauto.application.findwindows.ElementNotFoundError:
        print("Cửa sổ Word đã được đóng trước đó hoặc không tìm thấy.")
    except Exception as e:
        print(f"Lỗi khi đóng cửa sổ Word: {e}")
def check_doc_is_open(document_name):
    try:
        word_app = win32com.client.GetObject(Class="Word.Application")
    except win32com.client.com_error:
        # Word might not be running, return None
        return None
    for doc in word_app.Documents:
        if doc.Name == document_name:
            print("doc is : ",doc)
            return doc
        else:
            return word_app.ActiveDocument
def get_word_count(document_name=None):
    try:
        # Kết nối với ứng dụng Word đang chạy
        # word_app = win32com.client.Dispatch("Word.Application")
        word_app = win32com.client.GetObject(Class="Word.Application")
        # Kiểm tra xem có tài liệu nào đang mở không
        doc= check_doc_is_open(document_name)
        # Đếm số lượng từ
        word_count = doc.Words.Count
        return word_count
    except AttributeError as e:
        print(f"Lỗi thuộc tính: {e}")
        return None
    except Exception as e:
        print(f"Lỗi không xác định: {e}")
        return None


def wait_for_paste_completion(initial_word_count,document_name, timeout=30):
    start_time = time.time()
    while time.time() - start_time < timeout:
        current_word_count = get_word_count(document_name)
        if current_word_count > initial_word_count:
            return True
        time.sleep(1)
    return False
def find_element_with_timeout(driver, locator, timeout=60):
    """Tìm kiếm element với thời gian chờ tối đa.

    Args:
        driver: WebDriver instance.
        locator: Tuple (By.TYPE, "locator_string"). Ví dụ: (By.CLASS_NAME, "uk-margin-small-top").
        timeout: Thời gian chờ tối đa (giây). Mặc định là 60 giây.

    Returns:
        WebElement nếu tìm thấy, None nếu không tìm thấy sau thời gian chờ.
        Raises TimeoutException nếu quá thời gian chờ và vẫn không tìm thấy.
    """
    try:
        element = WebDriverWait(driver, timeout).until(EC.presence_of_element_located(locator))
        return element
    except TimeoutException:
        print(f"Không tìm thấy element với locator {locator} sau {timeout} giây.")
        return None  # hoặc raise lại TimeoutException nếu muốn dừng chương trình ngay lập tức
    except Exception as e: # Bắt các lỗi khác ngoài TimeoutException
        print(f"Lỗi không xác định: {e}")
        return None
def find_element_with_timeout(driver, locator, timeout=10): # Giảm timeout để test nhanh hơn
    try:
        element = WebDriverWait(driver, timeout).until(EC.presence_of_element_located(locator))
        return element
    except TimeoutException:
        return None  # Trả về None nếu không tìm thấy
    except Exception as e:
        print(f"Lỗi không xác định: {e}")
        return None

def find_element_by_multiple_locators(driver, locators):
    """Tìm element bằng cách thử nhiều locator theo thứ tự."""
    for locator in locators:
        element = find_element_with_timeout(driver, locator)
        if element:
            print(f"Đã tìm thấy element với locator: {locator}")
            return element
    return None  # Trả về None nếu không tìm thấy với bất kỳ locator nào
def process_element(driver, element, minimum_margin_top=30):
    if not element:
        print("Không có element nào để xử lý.")
        return

    driver.execute_script("arguments[0].scrollIntoView();", element)
    time.sleep(1)

    js_code = f"""
    var element = arguments[0];
    var minimumMarginTop = {minimum_margin_top};
    var stopElement = document.querySelector('h1, h2, h3[style*="margin-top: ' + minimumMarginTop + 'px"]');
    var range = document.createRange();

    range.setStartBefore(element);

    // SỬA LỖI Ở ĐÂY: Loại bỏ dấu ngoặc nhọn thừa
    if (stopElement) {
        range.setEndBefore(stopElement);
    } else {
        range.setEndAfter(document.body.lastChild);
    }

    var divs = document.querySelectorAll('[class*="uk-margin-remove-last-child"][class*="custom"]');
    divs.forEach(function(div) {
        if (div.querySelector('style')) {
            div.parentNode.removeChild(div);
        }
    });

    var sel = window.getSelection();
    sel.removeAllRanges();
    sel.addRange(range);
    """
    driver.execute_script(js_code, element)
   
def copy_and_paste_content(driver, document,output_path):
    locators = [
    (By.CLASS_NAME, "uk-margin-small-top"),
    (By.CSS_SELECTOR, ".uk-width-expand\\@m.uk-first-column"),
    (By.CSS_SELECTOR, ".some-other-class"), # Thêm các locator khác nếu cần
    (By.ID, "some-id") # Tìm kiếm theo ID
]

    try:
        element = find_element_by_multiple_locators(driver, locators)
        if element:
            process_element(driver, element)
        else:
            print("Không tìm thấy element nào với các locator đã cho.")
    except Exception as e:
        print(f"Lỗi tổng quát: {e}")
    
    if not element: # Nếu element class trên không tìm thấy thì tìm element class khác
        try:
            element = find_element_with_timeout(driver, (By.CSS_SELECTOR, ".uk-width-expand\\@m.uk-first-column"))
            if element:
                print("Đã tìm thấy element với CSS selector '.uk-width-expand\\@m.uk-first-column'")
            else:
                print("Không tìm thấy element với CSS selector '.uk-width-expand\\@m.uk-first-column'")

        except Exception as e:
            print(f"Lỗi tổng quát: {e}")
    time.sleep(1)
    pyautogui.hotkey('ctrl', 'c')
    path1 = "C:\\Program Files (x86)\\Microsoft Office\\root\\Office16\\winword.exe"
    path2 = "C:\\Program Files\\Microsoft Office\\root\\Office16\\winword.exe"
    if os.path.exists(path1):
        app_path = path1
    elif os.path.exists(path2):
        app_path = path2
    else:
        raise FileNotFoundError("Không tìm thấy Microsoft Word ở bất kỳ đường dẫn nào.")
    app = pywinauto.Application().start(app_path)
    
    # # Khởi động ứng dụng Word (nếu chưa mở)
    # try:
    #     app = Application(backend="uia").connect(title_re=".*Word.*", timeout=5) # Thay đổi backend nếu cần
    # except:
    #     app = Application(backend="uia").start("WINWORD.EXE") # Thay đổi backend nếu cần

    if not app.windows():
        app.start(app_path)

    try:
        dlg = app.window(title_re=".*Word.*")
        dlg.wait('visible', timeout=40)
        initial_screenshot = capture_screenshot(dlg)
        # Thử nhấn Ctrl+O và kiểm tra sự thay đổi của cửa sổ
        max_attempts = 5
        attempt = 0
        while attempt < max_attempts:
            dlg.type_keys('^o')
            time.sleep(2)  # Chờ một chút để cửa sổ có thể thay đổi
            if has_window_changed(dlg, initial_screenshot):
                print("Cửa sổ đã thay đổi.")
                break
            attempt += 1
            print(f"Thử lần {attempt} không thành công, thử lại...")
        if attempt == max_attempts:
            print("Mở file chờ 15s chưa có tín hiệu.")
        else:
            # Tiếp tục các lệnh khác sau khi cửa sổ thay đổi
            pass
        
        image_path = 'Python Tutorial\\browse_button_image.png'
        if not os.path.exists(image_path):
            print(f"Không tìm thấy file ảnh tại {image_path}. Vui lòng chọn file ảnh mới.")
            image_path = select_image_file(image_path)
        if not image_path:
            print("Không có file ảnh nào được chọn.")
        for i in range(10):
            browse_button_location = pyautogui.locateCenterOnScreen(image_path, confidence=0.8)
            if browse_button_location:
                break
        if browse_button_location:
            print(f"Đã tìm thấy file ảnh tại {image_path}.")
        else:
            print(f"Không tìm thấy file ảnh tại {image_path}. Vui lòng chọn file ảnh mới.")
            new_image_path = select_image_file(image_path)
            if new_image_path:
                imagePath = new_image_path
                browse_button_location = pyautogui.locateCenterOnScreen(imagePath)
                if browse_button_location:
                    print(f"Đã tìm thấy file ảnh tại {imagePath}.")
                else:
                    print("Không tìm thấy nút 'Browse' trong file ảnh mới.")
            else:
                print("Không có file ảnh nào được chọn.")
        print(f"Moving to: {browse_button_location}")
        pyautogui.moveTo(browse_button_location)
        time.sleep(2)
        pyautogui.click(browse_button_location)
        # user_dir = Path("C:/Users/12953 bao/Desktop/desktop/work/Project/Python/BasicLearnPython/W3schools")
        # output_path = user_dir / "output.docx"
        print(output_path)
        
        directory_path = os.path.dirname(output_path)
        file_name = os.path.basename(output_path)
        new_patch = Path(directory_path) / file_name
        
        # print("file name là : ", file_name)
        
        dlg_open = app.window(title_re=".*Open.*")
        dlg_open.wait('ready', timeout=3)
        dlg_open.type_keys(str(new_patch), with_spaces=True, pause=0.1)
        time.sleep(2)
        dlg_open.type_keys('{ENTER}')
        time.sleep(2)
        
        find_doc_name = ".*" + file_name+ ".*"
        print("find_doc_name là :", find_doc_name)
        
        # dlg = app.window(title_re=".*output.docx.*")
        # Chờ cửa sổ Word chính xuất hiện
        dlg_word_open = None
        try:
            # Sử dụng title_re phù hợp với ứng dụng Word của bạn. Ví dụ: ".*Document.* - Word"
            dlg_word_open = app.window(title_re=find_doc_name)
            dlg_word_open.wait('ready', timeout=10) # Đảm bảo cửa sổ word đã sẵn sàng
            print("File Word đã mở xong.")
            time.sleep(2)
            dlg_word_open.type_keys('^{END}')
            dlg_word_open.type_t.docxkeys('{ENTER}')
        except pywinauto.findwindows.ElementNotFoundError:
            print("Không tìm thấy cửa sổ Word. Có thể có lỗi khi mở file.")
        except Exception as e:
            print(f"Lỗi: {e}")        
        # Lấy số lượng từ ban đầu để kiểm tra sau khi dán
        initial_word_count = get_word_count(file_name)
        
        print(f"so ky tu den duoc là initial_word_count := '{initial_word_count}'")
        print("so ky tu den duoc là initial_word_count := ",initial_word_count)
        
        doc_check = check_doc_is_open(output_path)
        words_before=0
        words_after=0
        # Lấy số lượng từ trước khi dán
        if doc_check:
            # Use doc_check for further processing (assuming it's the correct document)
            try:
                words_before = doc_check.Words.Count
                print(f"Number of words before: {words_before}")
            except Exception as e:
                print(f"An error occurred while counting words: {e}")
            else:
                print(f"Document '{output_path}' not found open in Word.")
           
        
        pyautogui.hotkey('ctrl', 'v')
        # print("đã dán Ctrl + V")
        # Chờ cho đến khi số lượng từ thay đổi
        start_time = time.time()  # Ghi lại thời điểm bắt đầu
        timeout = 5 * 60  # 5 phút tính bằng giây
        
        while True:
            try:
                words_after = doc_check.Words.Count
                print(f"Number of words after: {words_after}")
            except Exception as e:
                print(f"An error occurred while counting words: {e}")

            if words_after > words_before:
                break
            else:
                print("Không có sự thay đổi khi dán")
            elapsed_time = time.time() - start_time
            if elapsed_time > timeout:
                print(f"Đã quá thời gian chờ ({timeout} giây). Dừng kiểm tra.")
                break

            time.sleep(1)  # Chờ 1 giây trước khi kiểm tra lại


        # Chờ dán hoàn thành
        # if wait_for_paste_completion(initial_word_count,file_name, timeout=50):
        #     print("Dán dữ liệu hoàn thành.")
        # else:
        #     print("Dán dữ liệu không hoàn thành trong thời gian chờ.")
        # try:
        pyautogui.hotkey('ctrl', 's')
            
        # Chờ lưu hoàn thành
        if wait_for_save_completion(output_path, timeout=30):
            print("Lưu file hoàn thành.")
        else:
            print("Lưu file không hoàn thành trong thời gian chờ.")
        
        dlg_word_open = app.window(title_re=find_doc_name)
        # Đóng cửa sổ Word (sử dụng hàm riêng để xử lý lỗi)
        close_word_document(dlg_word_open)
    except pywinauto.findwindows.ElementNotFoundError:
        print("Không tìm thấy cửa sổ Word")

def get_absolute_url(base_url, relative_url):
    return urljoin(base_url, relative_url)

def get_existing_links(document):
    existing_links = []
    for paragraph in document.paragraphs:
        if paragraph.style.name == 'Heading 2':
            existing_links.append(paragraph.text)
    return existing_links

base_url = 'https://vi.extendoffice.com'

def create_word_document(url, document, output_path):
    response = requests.get(url)
    soup = BeautifulSoup(response.content, 'html.parser')
    links = soup.find('ul', id='ul-search').find_all('a')

    # Gỡ lỗi: In thứ tự các liên kết
    for i, link in enumerate(links):
        if i<10 :
            print(f"Link {i}: {link.text.strip()} - {link['href']}")

    existing_links = get_existing_links(document)
    with open('links.txt', 'w', encoding='utf-8') as file:
        for i, link in enumerate(links):
            file.write(f'Link {i}: {link.text.strip()} - {link["href"]}\n')
            text = link.text.strip()
            print(f"Link la : '{text}'")
            href = link['href']
            if text in existing_links:
                print(f"Link '{text}' đã tồn tại trong tài liệu.")
                continue
            absolute_url = get_absolute_url(base_url, href)
            print(f"Đang truy cập link {i}: {absolute_url}")
            driver = webdriver.Chrome()
            driver.get(absolute_url)
            copy_and_paste_content(driver, document,output_path)
            driver.quit()
            print(f"Đã hoàn thành link {i}")

def main():
    root = Tk()
    root.withdraw()
    root.attributes('-topmost', True)
    response = messagebox.askyesno("Chọn tùy chọn", "Bạn có muốn tạo file mới không? (Chọn 'No' để cập nhật file cũ)")
    root.destroy()

    if response:
        document = Document()
        document.add_paragraph('')  # Tạo một đoạn văn bản trống để dán nội dung
        current_dir = os.getcwd()
        print("current dir là : " , current_dir)
        output_path = 'output.docx'
        output_path = os.path.join(current_dir, output_path)
        document.save(output_path)
    else:
        word_file_path = select_word_file()
        if word_file_path:
            document = Document(word_file_path)
            output_path = word_file_path
        else:
            print("Không có file nào được chọn.")
            return

    url = "https://vi.extendoffice.com/documents/excel"
    create_word_document(url, document, output_path)

if __name__ == "__main__":
    main()
