import os
import time
import pyautogui
import pywinauto
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait 
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
from urllib.parse import urljoin
from bs4 import BeautifulSoup
import requests
from tkinter import Tk, filedialog, messagebox
from pathlib import Path
from pywinauto.findwindows import ElementNotFoundError
from PIL import ImageGrab, ImageChops
import pyperclip

user_dir = Path("C:/Users/12953 bao/Desktop/desktop/work/Project/Python/BasicLearnPython/W3schools")
output_path = user_dir / "output.docx"

def select_image_file(default_path):
    root = Tk()
    root.withdraw()
    root.attributes('-topmost', True)
    file_path = filedialog.askopenfilename(initialdir=os.path.dirname(default_path), title="Select Image File",
                                           filetypes=(("PNG files", "*.png"), ("All files", "*.*")))
    root.destroy()
    return file_path

def select_word_file():
    root = Tk()
    root.withdraw()
    root.attributes('-topmost', True)
    file_path = filedialog.askopenfilename(title="Select Word File",
                                           filetypes=(("Word files", "*.docx"), ("All files", "*.*")))
    root.destroy()
    return file_path

# Hàm chụp ảnh màn hình của cửa sổ
def capture_screenshot(dlg):
    rect = dlg.rectangle()
    screenshot = ImageGrab.grab(bbox=(rect.left, rect.top, rect.right, rect.bottom))
    return screenshot

# Hàm kiểm tra sự thay đổi của cửa sổ
def has_window_changed(dlg, initial_screenshot):
    current_screenshot = capture_screenshot(dlg)
    diff = ImageChops.difference(initial_screenshot, current_screenshot)
    return diff.getbbox() is not None

def copy_and_paste_content(driver, document):
    try:
    # Thử tìm phần tử với class "uk-margin-small-top"
        element = WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.CLASS_NAME, "uk-margin-small-top")))
    except:
        # Nếu không tìm thấy, chuyển sang phần tử với class "uk-width-expand@m uk-first-column"
        element = WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.CSS_SELECTOR , ".uk-width-expand\\@m.uk-first-column"))
)

    driver.execute_script("arguments[0].scrollIntoView();", element)
    time.sleep(1)
    driver.execute_script("""
    var element = arguments[0];
    var stopElement = document.querySelector('.uk-margin-remove-last-child.custom h3[style="margin-top: 30px;"]');
    var range = document.createRange();

    if (stopElement) {
        range.setStartBefore(element);
        range.setEndBefore(stopElement);
    } else {
        range.setStartBefore(element);
        range.setEndAfter(document.body.lastChild);
    }

    // Lấy tất cả các phần tử <div> có class 'uk-margin-remove-last-child custom'
    var divs = document.querySelectorAll('.uk-margin-remove-last-child.custom');

    divs.forEach(function(div) {
        // Kiểm tra nếu phần tử con tiếp theo là <style>
        if (div.querySelector('style')) {
            div.parentNode.removeChild(div);
        }
    });

    var sel = window.getSelection();
    sel.removeAllRanges();
    sel.addRange(range);
    """, element)
    time.sleep(2)
    pyautogui.hotkey('ctrl', 'c')
    time.sleep(2)

    path1 = "C:\\Program Files (x86)\\Microsoft Office\\root\\Office16\\winword.exe"
    path2 = "C:\\Program Files\\Microsoft Office\\root\\Office16\\winword.exe"
    if os.path.exists(path1):
        app_path = path1
    elif os.path.exists(path2):
        app_path = path2
    else:
        raise FileNotFoundError("Không tìm thấy Microsoft Word ở bất kỳ đường dẫn nào.")
    app = pywinauto.Application().start(app_path)
    if not app.windows():
        app.start(app_path)

    try:
        dlg = app.window(title_re=".*Word.*")
        dlg.wait('visible', timeout=40)
        dlg.type_keys('^o')
        time.sleep(15)
        image_path = 'Python Tutorial\\browse_button_image.png'
        if not os.path.exists(image_path):
            print(f"Không tìm thấy file ảnh tại {image_path}. Vui lòng chọn file ảnh mới.")
            image_path = select_image_file(image_path)
        if not image_path:
            print("Không có file ảnh nào được chọn.")
        for i in range(10):
            browse_button_location = pyautogui.locateCenterOnScreen(image_path, confidence=0.8)
            if browse_button_location:
                break
        if browse_button_location:
            print(f"Đã tìm thấy file ảnh tại {image_path}.")
        else:
            print(f"Không tìm thấy file ảnh tại {image_path}. Vui lòng chọn file ảnh mới.")
            new_image_path = select_image_file(image_path)
            if new_image_path:
                imagePath = new_image_path
                browse_button_location = pyautogui.locateCenterOnScreen(imagePath)
                if browse_button_location:
                    print(f"Đã tìm thấy file ảnh tại {imagePath}.")
                else:
                    print("Không tìm thấy nút 'Browse' trong file ảnh mới.")
            else:
                print("Không có file ảnh nào được chọn.")
        print(f"Moving to: {browse_button_location}")
        pyautogui.moveTo(browse_button_location)
        time.sleep(2)
        pyautogui.click(browse_button_location)
        # user_dir = Path("C:/Users/12953 bao/Desktop/desktop/work/Project/Python/BasicLearnPython/W3schools")
        # output_path = user_dir / "output.docx"
        print(output_path)
        dlg_open = app.window(title_re=".*Open.*")
        dlg_open.wait('ready', timeout=10)
        dlg_open.type_keys(str(output_path), with_spaces=True, pause=0.1)
        time.sleep(10)
        dlg_open.type_keys('{ENTER}')
        time.sleep(10)
        dlg = app.window(title_re=".*output.docx.*")
        dlg.type_keys('^{END}')
        dlg.type_keys('{ENTER}')
        pyautogui.hotkey('ctrl', 'v')
        time.sleep(10)
        try:
            pyautogui.hotkey('ctrl', 's')
            time.sleep(10)
        except pywinauto.application.ElementNotFoundError as e:
            print(f"Không tìm thấy menu: {e}")
        except Exception as e:
            print(f"Lỗi không xác định: {e}")
        finally:
            dlg.close()
    except pywinauto.findwindows.ElementNotFoundError:
        print("Không tìm thấy cửa sổ Word")

def get_absolute_url(base_url, relative_url):
    return urljoin(base_url, relative_url)

def get_existing_links(document):
    existing_links = []
    for paragraph in document.paragraphs:
        if paragraph.style.name == 'Heading 2':
            existing_links.append(paragraph.text)
    return existing_links

base_url = 'https://vi.extendoffice.com'

def create_word_document(url, document, output_path):
    response = requests.get(url)
    soup = BeautifulSoup(response.content, 'html.parser')
    links = soup.find('ul', id='ul-search').find_all('a')

    # Gỡ lỗi: In thứ tự các liên kết
    for i, link in enumerate(links):
        if i<10 :
            print(f"Link {i}: {link.text.strip()} - {link['href']}")

    existing_links = get_existing_links(document)
    with open('links.txt', 'w', encoding='utf-8') as file:
        for i, link in enumerate(links):
            file.write(f'Link {i}: {link.text.strip()} - {link["href"]}\n')
            text = link.text.strip()
            href = link['href']
            if text in existing_links:
                print(f"Link '{text}' đã tồn tại trong tài liệu.")
                continue
            absolute_url = get_absolute_url(base_url, href)
            print(f"Đang truy cập link {i}: {absolute_url}")
            driver = webdriver.Chrome()
            driver.get(absolute_url)
            copy_and_paste_content(driver, document)
            driver.quit()
            print(f"Đã hoàn thành link {i}")

def main():
    root = Tk()
    root.withdraw()
    root.attributes('-topmost', True)
    response = messagebox.askyesno("Chọn tùy chọn", "Bạn có muốn tạo file mới không? (Chọn 'No' để cập nhật file cũ)")
    root.destroy()

    if response:
        document = Document()
        document.add_paragraph('')  # Tạo một đoạn văn bản trống để dán nội dung
        output_path = 'output.docx'
        document.save(output_path)
    else:
        word_file_path = select_word_file()
        if word_file_path:
            document = Document(word_file_path)
            output_path = word_file_path
        else:
            print("Không có file nào được chọn.")
            return

    url = "https://vi.extendoffice.com/documents/excel"
    create_word_document(url, document, output_path)

if __name__ == "__main__":
    main()
